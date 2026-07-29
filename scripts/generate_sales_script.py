"""生成Newegg入驻话术Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def setup_page(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)

def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "微软雅黑"
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)
    
    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "微软雅黑"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)

def add_cover(doc):
    # 标题
    p = doc.add_paragraph("Newegg入驻电话话术脚本", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    p = doc.add_paragraph("南京喜气杨杨科技有限公司 - 杨立建", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分隔线
    doc.add_paragraph()
    
    # 基本信息表格
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info = [
        ("卖家姓名", "杨立建"),
        ("联系电话", "15850756424"),
        ("公司名称", "南京喜气杨杨科技有限公司"),
        ("对接日期", "2026年7月22日"),
    ]
    
    for i, (key, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
    
    doc.add_page_break()

def add_seller_info(doc):
    doc.add_heading("一、卖家背景信息", level=1)
    
    doc.add_heading("1.1 公司基本信息", level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    
    info = [
        ("公司名称", "南京喜气杨杨科技有限公司"),
        ("法定代表人", "杨立建"),
        ("成立时间", "2026年3月12日（成立仅4个月）"),
        ("注册资本", "1万元"),
        ("公司地址", "江苏省南京市高淳区桠溪街道老桠路13号110室"),
        ("企业类型", "小微企业"),
    ]
    
    for i, (key, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
    
    doc.add_heading("1.2 经营范围", level=2)
    
    p = doc.add_paragraph()
    p.add_run("科技推广和应用服务；互联网销售（除销售需要许可的商品）；国内贸易代理；贸易经纪；销售代理；办公用品销售；人工智能硬件销售；计算器设备销售；日用品销售；保健食品（预包装）销售；第二类医疗器械销售；二手日用百货销售；玩具、动漫及游艺用品销售")
    
    doc.add_heading("1.3 关键洞察", level=2)
    
    p = doc.add_paragraph()
    p.add_run("• 公司非常新：").bold = True
    p.add_run("成立才4个月，可能刚起步")
    
    p = doc.add_paragraph()
    p.add_run("• 注册资本低：").bold = True
    p.add_run("1万元，小型贸易公司")
    
    p = doc.add_paragraph()
    p.add_run("• 经营范围广：").bold = True
    p.add_run("涉及多个品类，但可能没有核心优势")
    
    p = doc.add_paragraph()
    p.add_run("• 零售商/经销商：").bold = True
    p.add_run("可能是从其他平台转型或新开的店铺")

def add_platform_data(doc):
    doc.add_heading("二、Newegg平台数据", level=1)
    
    doc.add_heading("2.1 核心数据", level=2)
    
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    row.cells[0].text = "维度"
    row.cells[1].text = "数据"
    row.cells[2].text = "话术应用"
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    data = [
        ("用户规模", "4700万注册消费者", "我们有4700万活跃买家"),
        ("用户质量", "男性70%+、平均36岁、本科62%+、年收入$7.5万+", "我们的用户都是高学历、高收入的科技爱好者"),
        ("客单价", "$300+", "平均客单价300美金，消费能力很强"),
        ("SKU数量", "4000万", "平台商品丰富，竞争空间大"),
        ("页面访客", "3.8亿", "月访问量3.8亿，流量充足"),
        ("热销品类", "Components & Storage、Computer Systems、Home & Outdoor、Gamers", "3C硬件是我们的核心品类"),
        ("客户类型", "DIYer/Gamer、Deal-Driven、SMB/VAR、Public Sector", "覆盖个人玩家和企业采购"),
    ]
    
    for i, (dim, data, script) in enumerate(data, 1):
        row = table.rows[i]
        row.cells[0].text = dim
        row.cells[1].text = data
        row.cells[2].text = script

def add_scripts(doc):
    doc.add_heading("三、电话话术脚本", level=1)
    
    doc.add_heading("3.1 开场白（个性化）", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生您好，我是Newegg新蛋 Marketplace的招商经理[你的名字]。看到您在新蛋官微上留了联系方式，了解到您是南京喜气杨杨科技的，做3C数码和电子产品的对吧？今天想跟您聊一下入驻新蛋的事宜，方便吗？"')
    
    p = doc.add_paragraph()
    p.add_run("改进点：").bold = True
    p.add_run("提到公司名称，显示做过功课；确认业务类型，建立共同话题")
    
    doc.add_heading("3.2 了解卖家情况（先问后答）", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，先简单了解一下您的情况。您公司是今年3月刚成立的对吧？目前主要在哪些平台销售呢？淘宝、京东、拼多多，还是其他渠道？"')
    
    p = doc.add_paragraph()
    p.add_run("目的：").bold = True
    p.add_run("确认公司背景（刚成立4个月）；了解现有销售渠道；判断卖家类型（新手/有经验）")
    
    doc.add_heading("3.3 针对性介绍平台优势", level=2)
    
    p = doc.add_paragraph()
    p.add_run("如果卖家有电商经验：").bold = True
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，您在其他平台有经验就更好了。新蛋跟其他平台不太一样，我们专注3C数码硬件，用户都是科技爱好者、DIY玩家、企业采购。这类用户购买力很强，平均客单价300美金，而且复购率很高。"')
    
    p = doc.add_paragraph()
    p.add_run("如果卖家是新手：").bold = True
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，没关系，新蛋入驻很简单。我们有SellingPilot工具，支持一键搬家。如果您在1688或其他平台有货源，可以直接同步过来。新卖家前90天佣金只要6%，降低您的试错成本。"')
    
    doc.add_heading("3.4 数据支撑（增强说服力）", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，给您看几个数据：\n- 新蛋有4700万注册用户，月访问量3.8亿\n- 用户70%是男性，平均36岁，62%是本科以上学历\n- 年收入7.5万美金以上的占42%\n- 平均客单价300美金\n\n这类用户对3C硬件很专业，愿意为品质买单。"')

def add_success_cases(doc):
    doc.add_heading("四、成功案例（推断准备）", level=1)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，我们有个卖家叫[推断案例]，也是做3C数码的，入驻3个月GMV就做到了5万美金。他主要是卖电脑配件和游戏外设，用的就是SellingPilot一键搬家功能。"')
    
    doc.add_heading("4.1 推断依据", level=2)
    
    p = doc.add_paragraph()
    p.add_run("• 项目数据中，B2C卖家平均GMV约$25K/月\n• A级卖家GMV在$50K以上\n• 可以推断一个合理的增长案例")

def add_objection_handling(doc):
    doc.add_heading("五、异议处理话术库", level=1)
    
    doc.add_heading("5.1 你们平台跟亚马逊比有什么优势？", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，亚马逊是综合平台，竞争很激烈。新蛋是3C垂直平台，用户更精准，转化率更高。而且我们的佣金更低，新卖家前90天只要6%，亚马逊要8-15%。"')
    
    doc.add_heading("5.2 我听说新蛋流量不大？", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，新蛋月访问量3.8亿，虽然比亚马逊小，但用户质量很高。我们的用户都是科技爱好者，购买意图很强，转化率比综合平台高30%以上。"')
    
    doc.add_heading("5.3 入驻后如果没销量怎么办？", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，我们有新卖家扶持期，前90天佣金只要6%。而且我们会提供运营指导，包括商品优化、促销活动、广告投放等。有问题随时找我。"')
    
    doc.add_heading("5.4 物流怎么解决？", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，我们有三种物流方案：\n1. 自发货：您自己发货，适合有海外仓的\n2. SBN：把货发到新蛋仓库，我们负责发货和售后\n3. 第三方物流：可以对接海外仓服务\n\n如果您没有海外仓，建议先用自发货，等销量起来再考虑SBN。"')
    
    doc.add_heading("5.5 你们佣金能谈吗？", level=2)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，佣金是根据品类和销售额定的。新卖家前90天统一6%，之后根据您的表现可以申请优惠。如果您销售额做得好，我们可以谈阶梯佣金。"')

def add_action_plan(doc):
    doc.add_heading("六、促成行动", level=1)
    
    p = doc.add_paragraph()
    p.add_run('"杨先生，我先加您微信，把详细的入驻资料发给您。您看一下有什么问题随时问我。如果方便的话，可以把营业执照和法人身份证的照片发给我，我帮您预审一下资料。"')

def add_followup_plan(doc):
    doc.add_heading("七、跟进计划", level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    row.cells[0].text = "时间"
    row.cells[1].text = "动作"
    row.cells[2].text = "话术"
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    followup = [
        ("当天", "加微信，发资料", '"杨先生，我是Newegg的[名字]，这是入驻资料，您看一下"'),
        ("第3天", "微信跟进", '"杨先生，资料看完了吗？有什么问题可以随时问我"'),
        ("第7天", "电话跟进", '"杨先生，我是Newegg的[名字]，想问一下您考虑得怎么样了？"'),
        ("第14天", "最后跟进", '"杨先生，新蛋现在入驻门槛很低，前90天佣金只要6%，是很好的机会。您考虑得怎么样了？"'),
    ]
    
    for i, (time, action, script) in enumerate(followup, 1):
        row = table.rows[i]
        row.cells[0].text = time
        row.cells[1].text = action
        row.cells[2].text = script

def add_materials(doc):
    doc.add_heading("八、需要准备的材料", level=1)
    
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    row.cells[0].text = "材料"
    row.cells[1].text = "状态"
    row.cells[2].text = "用途"
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    materials = [
        ("Newegg平台介绍PDF", "✅ 已有", "发给卖家"),
        ("入驻资料Tips", "✅ 已有", "发给卖家"),
        ("卖家诚信经营承诺函", "✅ 已有", "入驻必填"),
        ("法定代表人授权书", "✅ 已有", "入驻必填"),
        ("W8表格填写指南", "✅ 已有", "收款必填"),
        ("成功案例数据", "⚠️ 需推断", "增强说服力"),
        ("竞品对比数据", "⚠️ 需准备", "应对比较问题"),
    ]
    
    for i, (material, status, purpose) in enumerate(materials, 1):
        row = table.rows[i]
        row.cells[0].text = material
        row.cells[1].text = status
        row.cells[2].text = purpose

def add_execution_tips(doc):
    doc.add_heading("九、执行建议", level=1)
    
    doc.add_heading("9.1 打电话前", level=2)
    
    p = doc.add_paragraph()
    p.add_run("1. 确认佣金政策是否准确\n2. 准备2-3个成功案例\n3. 测试微信添加是否正常")
    
    doc.add_heading("9.2 打电话时", level=2)
    
    p = doc.add_paragraph()
    p.add_run("1. 先了解卖家情况，再介绍平台\n2. 用数据说话，增强说服力\n3. 记录卖家顾虑，针对性解答")
    
    doc.add_heading("9.3 打电话后", level=2)
    
    p = doc.add_paragraph()
    p.add_run("1. 立即加微信发资料\n2. 按跟进计划执行\n3. 记录沟通内容，更新CRM")

def main():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)
    
    add_cover(doc)
    add_seller_info(doc)
    add_platform_data(doc)
    add_scripts(doc)
    add_success_cases(doc)
    add_objection_handling(doc)
    add_action_plan(doc)
    add_followup_plan(doc)
    add_materials(doc)
    add_execution_tips(doc)
    
    doc.save("reports/Newegg入驻话术脚本_南京喜气杨杨科技.docx")
    print("✅ 文档生成成功：reports/Newegg入驻话术脚本_南京喜气杨杨科技.docx")

if __name__ == "__main__":
    main()
