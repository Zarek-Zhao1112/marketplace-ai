"""
Newegg AM新人手册 - Word文档生成器
用法: python scripts/generate_guide.py

依赖: python-docx (pip install python-docx)
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
MD_PATH = BASE_DIR / "docs" / "am_newcomer_guide.md"
OUTPUT_PATH = BASE_DIR / "docs" / "AM新人手册.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_style(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "FF4655")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    return table


def parse_md_to_docx(md_path, output_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    current_section = None
    current_subsection = None
    in_list = False
    in_table = False
    table_buffer = []
    is_after_table = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        is_after_table = False

        if not stripped:
            in_list = False
            continue

        if stripped.startswith("---"):
            continue

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            in_table = True
            continue
        else:
            if in_table and table_buffer:
                rows_data = []
                for tline in table_buffer:
                    parts = [p.strip() for p in tline.split("|")[1:-1]]
                    if not all(p.startswith("-") for p in parts):
                        rows_data.append(parts)
                if rows_data:
                    add_table_with_style(doc, rows_data[0], rows_data[1:])
                table_buffer = []
                in_table = False
                is_after_table = True
                doc.add_paragraph("")
                continue

        if stripped.startswith("#### "):
            text = stripped[5:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            continue

        if stripped.startswith("### "):
            text = stripped[4:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1A, 0x26, 0x34)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
            continue

        if stripped.startswith("## "):
            text = stripped[3:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0xFF, 0x46, 0x55)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(8)
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0xFF, 0x46, 0x55)
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(4)

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = subtitle.add_run("面向实习生 Account Manager 的 Seller Portal + BI 系统实操指南")
            run2.font.size = Pt(12)
            run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph("")
            continue

        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            ch = p.add_run("• ")
            ch.font.size = Pt(11)
            r = p.add_run(text)
            r.font.size = Pt(11)
            in_list = True
            continue

        if stripped.startswith("1. ") or stripped.startswith("2. ") or \
           stripped.startswith("3. ") or stripped.startswith("4. ") or \
           stripped.startswith("5. ") or stripped.startswith("6. ") or \
           stripped.startswith("7. ") or stripped.startswith("8. "):
            in_list = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(stripped)
            r.font.size = Pt(11)
            continue

        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        if "**" in stripped:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(11)
                elif part.strip():
                    r = p.add_run(part)
                    r.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        p = doc.add_paragraph()
        r = p.add_run(stripped)
        r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)
    return output_path


def main():
    print(f"[*] 读取: {MD_PATH}")
    if not MD_PATH.exists():
        print(f"[!] 未找到文件: {MD_PATH}")
        return

    try:
        import docx
    except ImportError:
        print("[!] 需要安装 python-docx: pip install python-docx")
        return

    output = parse_md_to_docx(MD_PATH, OUTPUT_PATH)
    print(f"[OK] 已生成: {output} ({output.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
