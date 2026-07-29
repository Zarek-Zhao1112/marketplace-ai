"""生成77个seller调研报告Word文档"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def setup_page(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)

def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "微软雅黑"
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)
    
    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "微软雅黑"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)

def add_cover(doc):
    p = doc.add_paragraph("Newegg卖家调研报告", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("77个卖家数据分析", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info = [
        ("报告日期", "2026年7月24日"),
        ("数据范围", "77个卖家"),
        ("数据来源", "SKU分析JSON文件"),
    ]
    
    for i, (key, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
    
    doc.add_page_break()

def add_summary(doc, report):
    doc.add_heading("一、数据概览", level=1)
    
    summary = report["summary"]
    
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    
    data = [
        ("总卖家数", f"{summary['total_sellers']}个"),
        ("总GMV", f"${summary['total_gmv']:,.0f}"),
        ("平均GMV", f"${summary['avg_gmv']:,.0f}"),
        ("数据周期", summary['date']),
    ]
    
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value

def add_site_distribution(doc, report):
    doc.add_heading("二、站点分布", level=1)
    
    site_stats = report["site_distribution"]
    
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    row.cells[0].text = "站点"
    row.cells[1].text = "卖家数"
    row.cells[2].text = "总GMV"
    row.cells[3].text = "占比"
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    total_gmv = sum(data["total_gmv"] for data in site_stats.values())
    
    for i, (site, data) in enumerate(site_stats.items(), 1):
        row = table.rows[i]
        row.cells[0].text = site
        row.cells[1].text = str(data["count"])
        row.cells[2].text = f"${data['total_gmv']:,.0f}"
        row.cells[3].text = f"{data['total_gmv']/total_gmv*100:.1f}%"

def add_grade_distribution(doc, report):
    doc.add_heading("三、等级分布", level=1)
    
    grade_stats = report["grade_distribution"]
    total = sum(grade_stats.values())
    
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    row.cells[0].text = "等级"
    row.cells[1].text = "卖家数"
    row.cells[2].text = "占比"
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for i, (grade, count) in enumerate(grade_stats.items(), 1):
        row = table.rows[i]
        row.cells[0].text = f"{grade}级"
        row.cells[1].text = str(count)
        row.cells[2].text = f"{count/total*100:.1f}%"

def add_top_sellers(doc, report):
    doc.add_heading("四、Top10卖家", level=1)
    
    top_sellers = report["top_sellers"]
    
    table = doc.add_table(rows=11, cols=6)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    headers = ["排名", "SellerID", "卖家名称", "站点", "GMV", "等级"]
    for j, header in enumerate(headers):
        row.cells[j].text = header
        for paragraph in row.cells[j].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for i, seller in enumerate(top_sellers, 1):
        row = table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = seller["sid"]
        row.cells[2].text = seller["name"][:20] if seller["name"] else "-"
        row.cells[3].text = seller["site"]
        row.cells[4].text = f"${seller['gmv']:,.0f}"
        row.cells[5].text = seller["grade"]

def add_category_analysis(doc, report):
    doc.add_heading("五、品类分析", level=1)
    
    category_analysis = report["category_analysis"]
    
    table = doc.add_table(rows=min(11, len(category_analysis) + 1), cols=4)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    headers = ["品类", "卖家数", "总GMV", "平均GMV"]
    for j, header in enumerate(headers):
        row.cells[j].text = header
        for paragraph in row.cells[j].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for i, (category, data) in enumerate(category_analysis[:10], 1):
        if i >= len(table.rows):
            break
        row = table.rows[i]
        row.cells[0].text = category[:20]
        row.cells[1].text = str(data["count"])
        row.cells[2].text = f"${data['total_gmv']:,.0f}"
        row.cells[3].text = f"${data['total_gmv']/data['count']:,.0f}"

def add_insights(doc, report):
    doc.add_heading("六、关键洞察", level=1)
    
    insights = report["insights"]
    
    for i, insight in enumerate(insights, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(insight)

def add_recommendations(doc):
    doc.add_heading("七、选品建议", level=1)
    
    recommendations = [
        ("高GMV品类", "优先选择GMV高的品类，如电脑配件、游戏外设", "GMV高意味着市场需求大"),
        ("低竞争品类", "关注竞争较小的细分品类", "新卖家更容易出单"),
        ("高毛利品类", "选择毛利率高的品类", "利润空间大，可持续发展"),
        ("季节性品类", "关注开学季、黑五、圣诞等节点", "抓住促销机会"),
    ]
    
    table = doc.add_table(rows=len(recommendations) + 1, cols=3)
    table.style = "Table Grid"
    
    # 表头
    row = table.rows[0]
    headers = ["建议类型", "具体内容", "理由"]
    for j, header in enumerate(headers):
        row.cells[j].text = header
        for paragraph in row.cells[j].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for i, (type_, content, reason) in enumerate(recommendations, 1):
        row = table.rows[i]
        row.cells[0].text = type_
        row.cells[1].text = content
        row.cells[2].text = reason

def add_action_items(doc):
    doc.add_heading("八、行动清单", level=1)
    
    actions = [
        "1. 了解目标卖家的背景和需求",
        "2. 根据卖家条件筛选适合的品类",
        "3. 分析目标品类的竞争情况",
        "4. 评估利润空间和风险",
        "5. 制定选品方案和上架计划",
        "6. 跟踪出单情况，及时调整策略",
    ]
    
    for action in actions:
        p = doc.add_paragraph(action)

def main():
    # 读取报告数据
    with open('data/seller_analysis_report.json', 'r', encoding='utf-8') as f:
        report = json.load(f)
    
    # 创建Word文档
    doc = Document()
    setup_page(doc)
    tune_styles(doc)
    
    add_cover(doc)
    add_summary(doc, report)
    add_site_distribution(doc, report)
    add_grade_distribution(doc, report)
    add_top_sellers(doc, report)
    add_category_analysis(doc, report)
    add_insights(doc, report)
    add_recommendations(doc)
    add_action_items(doc)
    
    # 保存文档
    doc.save("reports/77个seller调研报告.docx")
    print("✅ 文档生成成功：reports/77个seller调研报告.docx")

if __name__ == "__main__":
    main()
